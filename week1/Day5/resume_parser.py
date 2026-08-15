import os
import json
from pathlib import Path
import time

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

# Load environment variables
load_dotenv(Path(__file__).resolve().parents[2] / ".env")


# Get API key
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("api error")


# Register Groq client
client = Groq(api_key=my_api_key)

model = "llama-3.3-70b-versatile"

print("\n Resume Evaluator \n")

jd_path = Path("job_description.txt")


if not jd_path.exists():
    raise ValueError("Job description file not found.")

job_description = jd_path.read_text(encoding="utf-8")

if not job_description.strip():
    raise ValueError("Job description cannot be empty.")

class JobDescription(BaseModel):
    role : str
    required_skills : list[str]
    preferred_skills : list[str]
    minimum_experience : float | None
    educational_requirements : list[str]
    responsibilities : list[str]

job_schema = JobDescription.model_json_schema()


system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:
{job_schema}

IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information. 
"""

user_prompt = f"""
Analyse the following job description:
{job_description}
"""

system_message = {
    "role" : "system",
    "content" : system_prompt
}

user_message = {
    "role" : "user",
    "content" : user_prompt
}

response_format = { 
    "type" : "json_object"
}

messages = [system_message, user_message]

response = client.chat.completions.create(model = model, messages = messages, response_format = response_format)

answer = response.choices[0].message.content

raw_json = answer 
# print(raw_json)

job_data = json.loads(raw_json)

job = JobDescription(**job_data)

print(job.minimum_experience)
print(job.educational_requirements)

class MatchDetails(BaseModel):
    candidate_name: str
    matching_skills: list[str]
    missing_important_skills: list[str]
    experience_requirement_met: bool
    final_verdict: str

class MatchResult(BaseModel):
    score : float
    details : MatchDetails

class  Experience(BaseModel):
    company : str | None = None
    role : str | None = None
    duration : str | None = None
    description : str | None = None
    skills_used : list[str] = Field(default_factory = list)

class Resume(BaseModel):
    name: str | None = None
    email : str | None = None
    phone : str | None = None

    total_experience_years : float | None  =  None

    skills : list[str] = Field(default_factory = list)
    experiences : list[Experience] = Field(default_factory = list)
    education : list[str] = Field(default_factory = list)
    projects : list[str] = Field(default_factory = list)
    certifications : list[str] = Field(default_factory = list)

resume_schema = Resume.model_json_schema()

def final_score(job, resume):
    match_schema =  MatchResult.model_json_schema()
    prompt = f""" 
    You're an HR recruiter.
     
    Compare the candidates resume with job description.
    
     Job Description: 
    {job.model_dump_json(indent = 2)}

    Candidate Resume:
    {resume.model_dump_json(indent = 2)}
    Return json matching this schema:
    {match_schema}
    
    Give me:
    1. Candidate Nmae
    2. Matching Skills
    3.Missing Important Skils
    4. Whether experience requirement is met.
    5. Overall match percentage from 0 to 100
    6. A short final verdict.

    Keep the response easy to read and concise.
    """
    message = {
        "role" : "user",
        "content" : prompt
    }

    messages = [message] 
    response_format = {
        "type" : "json_object"
    }

    response =  client.chat.completions.create(model = model, messages = messages, response_format = response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an experet resume parser.

    Extract information from the resume based on the meaining,
    not only based on the exact section headings.

    Different resumes may use different headings.

    for example: 
    - experience
    - professional experience
    - work history
    -employment
    -internships

    They all contain relaant experience

    skills may also appear in the skills section, work expereince, internships or projects

    return only valid json matching this schema:
    {resume_schema}

    Imprtant rules:

    1. Do Not invent information.
    2. If value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """

    user_prompt = f"""
    Parse the following resume:

        {resume_text}
        """

    message_system = {
        "role" : "system",
        "content" : system_prompt
    }

    message_user = {
        "role" : "user",
        "content" : user_prompt
    }

    messages = [message_system, message_user]

    response_format = {
        "type" : "json_object"
    }

    response = client.chat.completions.create(model = model, messages = messages, response_format = response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)


    if data.get("skills") is None:
        data["skills"] = []

    if data.get("education") is None:
        data["education"] = []

    if data.get("projects") is None:
        data["projects"] = []

    if data.get("certifications") is None:
        data["certifications"] = []

    if data.get("experiences") is None:
        data["experiences"] = []

    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document

# function to read pdf file
def read_pdf(file_path):
    try:
        reader = PdfReader(file_path)

        text = ""

        for page in reader.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

        return text

    except Exception as e:
        raise ValueError(f"Could not read PDF: {e}")


# function to read docx file
def read_docx(file_path):
    try:
        document = Document(file_path)

        text = ""

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        return text

    except Exception as e:
        raise ValueError(f"Could not read DOCX: {e}")
    
# function to read resume 
def read_resume(file_path):
    file_path = Path(file_path)
    if file_path.suffix.lower() == ".pdf" :
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only pdf and docx files are supported")

resume_folder = Path("resumes") 
all_results = []

for file_path in resume_folder.iterdir():
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue

    print("\n Processing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)
    time.sleep(5)
    # rate limiting to prevent dos attack   
    result = final_score(job, parsed_resume)
    time.sleep(5)
    print("Score:",  result.score)
    all_results.append({
        "name" : parsed_resume.name,
        "score" : result.score,
        "details": result.details
    })

    all_results.sort(
        key = lambda candidate : candidate["score"],
        reverse = True
    )

top_2 = all_results[:2]
worst_2 = all_results[-2:]

print("TOP  2 Candidates")

for candidate in top_2:
    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

    print("Lowest 2 candidates")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])
